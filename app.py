import io
import os
import docx
import glob
import flask
import webvtt
import zipfile
import tempfile
import traceback
import youtube_dl

app = flask.Flask(__name__)

@app.errorhandler(Exception)
def handle_bad_request(e):
    return 'Exception occured:\n{}'.format(traceback.format_exc()), \
            400, {'Content-Type': 'text/plain'}

@app.route("/")
def home():
    return flask.render_template('home.html')

@app.route('/ytdl')
def download_files():
    id = flask.request.args.get('id')
    if id.startswith('http'):
        url = id
        id = url[-11:]
    else:
        url = "https://www.youtube.com/watch?v={}".format(id)

    with tempfile.TemporaryDirectory() as directory:
        os.chdir(directory)

        # youtube-dl
        ydl_opts = {
            'writethumbnail': True,
            'writeautomaticsub': True,
            'subtitleslangs': ['en', 'es'],
            'skip_download': True,
        }

        with youtube_dl.YoutubeDL(ydl_opts) as ydl:
            ydl.extract_info(url, download=True)

        # webvtt & docx
        if flask.request.args.get('concat') != None:
            document = docx.Document()
            table = document.add_table(rows=1, cols=2)
            cells = table.rows[0].cells

            num = 0
            fname = ""
            for file in glob.glob("*.vtt"):

                vtt = webvtt.read(file)
                transcript = ""

                lines = []
                for line in vtt:
                    lines.extend(line.text.strip().splitlines())

                prev = None
                for line in lines:
                    if line == prev:
                        continue
                    transcript += line + "\r"
                    prev = line

                cells[num].text = transcript
                fname = file
                num += 1

            document.save("{}.docx".format(fname[:-7]))
        
        else:
            for file in glob.glob("*.vtt"):
                document = docx.Document()

                vtt = webvtt.read(file)
                transcript = ""

                lines = []
                for line in vtt:
                    lines.extend(line.text.strip().splitlines())

                prev = None
                for line in lines:
                    if line == prev:
                        continue
                    transcript += line + "\r"
                    prev = line

                document.add_paragraph(transcript)
                document.save("{}.docx".format(file))

        #zipfile
        files = set(glob.glob("*")) - set(glob.glob("*.vtt"))

        with io.BytesIO() as zip_buffer:
            with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED, False) as zip_file:
                for file in files:
                    zip_file.write(file)
            zip_buffer.seek(0)
            content = zip_buffer.getvalue()

        resp = flask.Response(content)
        filename = '{}.zip'.format(id)
        resp.headers['Content-Type'] = 'application/zip'
        resp.headers['Content-Disposition'] = 'inline; filename="{}"'.format(filename)
        resp.headers['Content-Length'] = len(content)

        return resp

if __name__ == "__main__":
    app.run(host='0.0.0.0')
